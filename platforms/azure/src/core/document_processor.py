"""
================================================================================
document_processor.py - 文書処理モジュール
================================================================================

【概要】
様々なファイル形式（PDF、Excel、テキスト等）からテキストを抽出します。
Azure Document Intelligence と連携することで、高精度なOCR処理も可能です。

【対応ファイル形式】
- テキスト系: .txt, .csv, .json, .xml, .log, .md
- PDF: .pdf（通常PDF / スキャンPDF）
- Excel: .xlsx, .xls
- 画像: .jpg, .jpeg, .png, .gif, .bmp, .webp, .tiff, .tif

【主要クラス】
- DocumentProcessor: テキスト抽出の統合インターフェース
- AzureDocumentIntelligence: Azure Document Intelligence クライアント
- ExtractedContent: 抽出結果を格納するデータクラス
- TextElement: テキスト要素（座標情報付き）
- ExtractedTable: 抽出された表データ

【Azure Document Intelligence について】
環境変数で設定すると、PDFや画像ファイルからOCRでテキスト抽出が可能です。
- AZURE_DI_ENDPOINT: Document Intelligence のエンドポイント
- AZURE_DI_KEY: APIキー

設定がない場合は、ローカルライブラリ（pypdf, openpyxl）にフォールバックします。

【使用例】
```python
from core.document_processor import DocumentProcessor

# 単一ファイルの抽出
result = DocumentProcessor.extract_text(
    file_name="report.pdf",
    extension=".pdf",
    base64_content="JVBERi0xLj..."
)
print(result.text_content)

# 複数ファイルの一括抽出
results = DocumentProcessor.extract_all(evidence_files)

# LLM プロンプト用にフォーマット
prompt_text = DocumentProcessor.format_for_prompt(results)
```

================================================================================
"""
import base64
import io
import os
import logging
from typing import Optional, List, Dict, Any
from dataclasses import dataclass, field

# ロガーの設定
logger = logging.getLogger(__name__)


# =============================================================================
# データクラス定義
# =============================================================================

@dataclass
class TextElement:
    """
    テキスト要素（座標情報付き）

    Document Intelligence で抽出した際、各テキスト要素の位置情報を保持します。
    将来的にPDFハイライト機能で使用する予定です。

    Attributes:
        element_id (str): 要素の一意識別子（例: "elem_0", "para_1"）
        text (str): テキスト内容
        page_number (int): ページ番号（1始まり）
        bounding_box (List[float]): 座標 [x1, y1, x2, y2]（インチ単位）
        element_type (str): 要素タイプ
            - "paragraph": 段落
            - "table_cell": 表のセル
            - "line": 行
            - "word": 単語
        confidence (float): 認識信頼度（0.0〜1.0）
    """
    element_id: str
    text: str
    page_number: int
    bounding_box: List[float]  # [x1, y1, x2, y2] インチ単位
    element_type: str  # "paragraph", "table_cell", "line", "word"
    confidence: float = 1.0


@dataclass
class TableCell:
    """
    表のセル情報

    表の各セルの内容と位置情報を保持します。

    Attributes:
        row_index (int): 行番号（0始まり）
        column_index (int): 列番号（0始まり）
        text (str): セル内のテキスト
        bounding_box (List[float]): 座標 [x1, y1, x2, y2]
        row_span (int): 行の結合数（デフォルト: 1）
        column_span (int): 列の結合数（デフォルト: 1）
    """
    row_index: int
    column_index: int
    text: str
    bounding_box: List[float]
    row_span: int = 1
    column_span: int = 1


@dataclass
class ExtractedTable:
    """
    抽出された表データ

    PDFやExcelから抽出した表の構造情報を保持します。

    Attributes:
        table_id (str): 表の識別子（例: "table_0"）
        page_number (int): ページ番号
        row_count (int): 行数
        column_count (int): 列数
        cells (List[TableCell]): セルの一覧
        bounding_box (List[float]): 表全体の座標
    """
    table_id: str
    page_number: int
    row_count: int
    column_count: int
    cells: List[TableCell]
    bounding_box: List[float]


@dataclass
class ExtractedContent:
    """
    文書から抽出した内容

    ファイルからテキストを抽出した結果を格納します。
    エラー情報や座標情報（Document Intelligence使用時）も含みます。

    Attributes:
        file_name (str): 元のファイル名
        file_type (str): ファイルタイプ（"pdf", "excel", "text"等）
        text_content (str): 抽出されたテキスト内容
        extraction_method (str): 抽出方法
            - "pypdf": pypdfライブラリ使用
            - "openpyxl": openpyxlライブラリ使用
            - "direct_decode": テキストファイル直接読み込み
            - "azure_document_intelligence_layout": Azure DI使用
            - "error": エラー発生
        page_count (int, optional): ページ数
        error (str, optional): エラーメッセージ
        elements (List[TextElement], optional): テキスト要素（座標情報付き）
        tables (List[ExtractedTable], optional): 抽出された表

    Note:
        elements と tables は Document Intelligence 使用時のみ設定されます。
    """
    file_name: str
    file_type: str
    text_content: str
    extraction_method: str
    page_count: Optional[int] = None
    error: Optional[str] = None
    # Document Intelligence 使用時の座標情報
    elements: Optional[List[TextElement]] = None
    tables: Optional[List[ExtractedTable]] = None


# =============================================================================
# Azure Document Intelligence クライアント
# =============================================================================

class AzureDocumentIntelligence:
    """
    Azure Document Intelligence クライアント

    高精度なOCR・レイアウト解析を提供します。
    環境変数で設定された場合に自動的に有効になります。

    環境変数:
        AZURE_DI_ENDPOINT: Document Intelligence のエンドポイントURL
        AZURE_DI_KEY: APIキー

    Example:
        >>> di = AzureDocumentIntelligence()
        >>> if di.is_configured():
        ...     result = di.extract_with_layout("report.pdf", base64_content)
        ...     print(result.text_content)

    Note:
        - azure-ai-documentintelligence パッケージが必要です
        - 従量課金の Azure サービスのため、コスト管理に注意
    """

    def __init__(self):
        """
        Document Intelligence クライアントを初期化

        環境変数からエンドポイントとAPIキーを取得します。
        """
        # 環境変数から設定を読み込み
        self.endpoint = os.getenv("AZURE_DI_ENDPOINT")
        self.key = os.getenv("AZURE_DI_KEY")
        self._client = None

        # 設定状態をログ出力
        if self.is_configured():
            logger.info("Azure Document Intelligence: 設定済み")
        else:
            logger.debug("Azure Document Intelligence: 未設定（ローカル抽出を使用）")

    def is_configured(self) -> bool:
        """
        Document Intelligence が設定されているかチェック

        Returns:
            bool: エンドポイントとキーが両方設定されていれば True
        """
        return bool(self.endpoint and self.key)

    def _get_client(self):
        """
        Document Intelligence クライアントを取得（遅延初期化）

        初回呼び出し時にクライアントを作成し、以降はキャッシュを返します。

        Returns:
            DocumentIntelligenceClient: Azure SDK クライアント

        Raises:
            ImportError: azure-ai-documentintelligence パッケージ未インストール
        """
        if self._client is None:
            try:
                # Azure SDK のインポート
                from azure.ai.documentintelligence import DocumentIntelligenceClient
                from azure.core.credentials import AzureKeyCredential

                # クライアント作成
                self._client = DocumentIntelligenceClient(
                    endpoint=self.endpoint,
                    credential=AzureKeyCredential(self.key)
                )
                logger.debug("Document Intelligence クライアントを作成しました")

            except ImportError:
                logger.error("azure-ai-documentintelligence パッケージがインストールされていません")
                raise ImportError(
                    "azure-ai-documentintelligence パッケージをインストールしてください: "
                    "pip install azure-ai-documentintelligence"
                )

        return self._client

    def extract_with_layout(self, file_name: str, base64_content: str,
                            mime_type: str = None) -> ExtractedContent:
        """
        Document Intelligence Layout API でテキストを抽出

        OCRによる文字認識とレイアウト解析を行い、座標情報付きで結果を返します。

        Args:
            file_name (str): ファイル名
            base64_content (str): Base64エンコードされたファイル内容
            mime_type (str, optional): MIMEタイプ

        Returns:
            ExtractedContent: 抽出結果（テキスト、座標情報、表データ含む）

        Note:
            - 処理時間はファイルサイズ・ページ数に依存します
            - 大きなファイルは数秒〜数十秒かかる場合があります
        """
        if not self.is_configured():
            raise ValueError("Document Intelligence が設定されていません")

        logger.info(f"[Document Intelligence] 抽出開始: {file_name}")

        try:
            from azure.ai.documentintelligence.models import AnalyzeDocumentRequest

            client = self._get_client()

            # Base64デコード
            file_bytes = base64.b64decode(base64_content)
            logger.debug(f"ファイルサイズ: {len(file_bytes):,} バイト")

            # Document Intelligence で解析
            logger.debug("Layout モデルで解析中...")
            poller = client.begin_analyze_document(
                model_id="prebuilt-layout",
                body=AnalyzeDocumentRequest(bytes_source=file_bytes),
                content_type=mime_type or "application/octet-stream"
            )

            result = poller.result()

            # 結果を処理
            text_parts = []
            elements = []
            tables = []
            element_counter = 0

            # ページ数を取得
            page_count = len(result.pages) if result.pages else 0
            logger.info(f"[Document Intelligence] ページ数: {page_count}")

            # 各ページを処理
            for page in result.pages or []:
                page_num = page.page_number
                text_parts.append(f"--- ページ {page_num} ---")

                # 行を処理
                for line in page.lines or []:
                    text_parts.append(line.content)

                    # 座標情報を保存
                    if line.polygon:
                        bbox = self._polygon_to_bbox(line.polygon)
                        elements.append(TextElement(
                            element_id=f"elem_{element_counter}",
                            text=line.content,
                            page_number=page_num,
                            bounding_box=bbox,
                            element_type="line",
                            confidence=1.0
                        ))
                        element_counter += 1

            # 段落を処理
            for para in result.paragraphs or []:
                if para.bounding_regions:
                    for region in para.bounding_regions:
                        bbox = self._polygon_to_bbox(region.polygon) if region.polygon else [0, 0, 0, 0]
                        elements.append(TextElement(
                            element_id=f"para_{element_counter}",
                            text=para.content,
                            page_number=region.page_number,
                            bounding_box=bbox,
                            element_type="paragraph",
                            confidence=1.0
                        ))
                        element_counter += 1

            # 表を処理
            for table_idx, table in enumerate(result.tables or []):
                table_cells = []
                table_bbox = [0, 0, 0, 0]

                for cell in table.cells or []:
                    cell_bbox = [0, 0, 0, 0]
                    cell_page = 1

                    if cell.bounding_regions:
                        region = cell.bounding_regions[0]
                        cell_page = region.page_number
                        if region.polygon:
                            cell_bbox = self._polygon_to_bbox(region.polygon)

                    table_cells.append(TableCell(
                        row_index=cell.row_index,
                        column_index=cell.column_index,
                        text=cell.content or "",
                        bounding_box=cell_bbox,
                        row_span=cell.row_span or 1,
                        column_span=cell.column_span or 1
                    ))

                    # セルも要素として保存
                    elements.append(TextElement(
                        element_id=f"cell_{table_idx}_{cell.row_index}_{cell.column_index}",
                        text=cell.content or "",
                        page_number=cell_page,
                        bounding_box=cell_bbox,
                        element_type="table_cell",
                        confidence=1.0
                    ))

                # 表の座標を取得
                if table.bounding_regions:
                    region = table.bounding_regions[0]
                    if region.polygon:
                        table_bbox = self._polygon_to_bbox(region.polygon)
                    table_page = region.page_number
                else:
                    table_page = 1

                tables.append(ExtractedTable(
                    table_id=f"table_{table_idx}",
                    page_number=table_page,
                    row_count=table.row_count or 0,
                    column_count=table.column_count or 0,
                    cells=table_cells,
                    bounding_box=table_bbox
                ))

            logger.info(f"[Document Intelligence] 表: {len(tables)}個")

            # テキストを結合
            full_text = "\n".join(text_parts)

            # 表データをテキスト形式で追加
            if tables:
                full_text += "\n\n=== 表データ ===\n"
                for table in tables:
                    full_text += f"\n[表 {table.table_id} ({table.row_count}行×{table.column_count}列)]\n"

                    # 表をテキスト形式に変換
                    table_rows: Dict[int, Dict[int, str]] = {}
                    for cell in table.cells:
                        if cell.row_index not in table_rows:
                            table_rows[cell.row_index] = {}
                        table_rows[cell.row_index][cell.column_index] = cell.text

                    for row_idx in sorted(table_rows.keys()):
                        row_data = table_rows[row_idx]
                        row_text = "\t".join(
                            row_data.get(col_idx, "")
                            for col_idx in range(table.column_count)
                        )
                        full_text += row_text + "\n"

            logger.info(f"[Document Intelligence] 抽出完了: {len(full_text):,}文字")

            return ExtractedContent(
                file_name=file_name,
                file_type="pdf" if file_name.lower().endswith('.pdf') else "image",
                text_content=full_text,
                extraction_method="azure_document_intelligence_layout",
                page_count=page_count,
                elements=elements,
                tables=tables
            )

        except ImportError as e:
            logger.error(f"Document Intelligence SDK が利用できません: {e}")
            raise

        except Exception as e:
            logger.error(f"Document Intelligence 抽出エラー ({file_name}): {e}")
            return ExtractedContent(
                file_name=file_name,
                file_type="unknown",
                text_content=f"[Document Intelligence エラー: {file_name}] - {str(e)}",
                extraction_method="azure_di_error",
                error=str(e)
            )

    def _polygon_to_bbox(self, polygon: List[float]) -> List[float]:
        """
        ポリゴン座標をバウンディングボックスに変換

        Document Intelligence の座標形式（8点）を [x1, y1, x2, y2] に変換します。

        Args:
            polygon: [x1, y1, x2, y2, x3, y3, x4, y4] 形式の座標

        Returns:
            [x1, y1, x2, y2] 形式のバウンディングボックス
        """
        if not polygon or len(polygon) < 4:
            return [0, 0, 0, 0]

        # X座標とY座標を分離
        x_coords = [polygon[i] for i in range(0, len(polygon), 2)]
        y_coords = [polygon[i] for i in range(1, len(polygon), 2)]

        return [
            min(x_coords),  # x1 (左端)
            min(y_coords),  # y1 (上端)
            max(x_coords),  # x2 (右端)
            max(y_coords)   # y2 (下端)
        ]


# =============================================================================
# メイン処理クラス
# =============================================================================

class DocumentProcessor:
    """
    文書テキスト抽出の統合インターフェース

    様々なファイル形式からテキストを抽出します。
    Azure Document Intelligence が設定されている場合は優先的に使用します。

    対応ファイル形式:
        - テキスト: .txt, .csv, .json, .xml, .log, .md
        - PDF: .pdf（通常 / スキャン）
        - Excel: .xlsx, .xls
        - 画像: .jpg, .jpeg, .png, .gif, .bmp, .webp, .tiff, .tif

    Class Methods:
        extract_text: 単一ファイルからテキスト抽出
        extract_all: 複数ファイルを一括抽出
        format_for_prompt: LLMプロンプト用にフォーマット
        get_config_status: 設定状態を取得

    Example:
        >>> # 単一ファイル抽出
        >>> result = DocumentProcessor.extract_text(
        ...     file_name="report.pdf",
        ...     extension=".pdf",
        ...     base64_content=content
        ... )
        >>>
        >>> # 複数ファイル一括抽出
        >>> results = DocumentProcessor.extract_all(evidence_files)
        >>>
        >>> # プロンプト用フォーマット
        >>> text = DocumentProcessor.format_for_prompt(results)
    """

    # 対応ファイル形式の定義
    TEXT_TYPES = ['.txt', '.csv', '.json', '.xml', '.log', '.md']
    PDF_TYPES = ['.pdf']
    EXCEL_TYPES = ['.xlsx', '.xls']
    WORD_TYPES = ['.docx', '.doc']
    IMAGE_TYPES = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.tiff', '.tif']

    # Document Intelligence クライアント（シングルトン）- 後方互換用
    _di_client: Optional[AzureDocumentIntelligence] = None

    @classmethod
    def get_di_client(cls) -> AzureDocumentIntelligence:
        """
        Document Intelligence クライアントを取得（シングルトン）

        後方互換のため残しています。新規コードはOCRFactoryを使用してください。

        Returns:
            AzureDocumentIntelligence: クライアントインスタンス
        """
        if cls._di_client is None:
            cls._di_client = AzureDocumentIntelligence()
        return cls._di_client

    @classmethod
    def get_ocr_client(cls):
        """
        OCRクライアントを取得（OCRFactoryを使用）

        OCR_PROVIDER環境変数に基づいて適切なOCRクライアントを返します。

        Returns:
            BaseOCRClient: OCRクライアント（NONE指定時はNone）
        """
        try:
            from infrastructure.ocr_factory import OCRFactory
            return OCRFactory.get_ocr_client()
        except ImportError:
            logger.warning("OCRFactory が利用できません。AzureDocumentIntelligenceにフォールバック")
            di_client = cls.get_di_client()
            return di_client if di_client.is_configured() else None

    @classmethod
    def extract_text(cls, file_name: str, extension: str, base64_content: str,
                     mime_type: str = None, use_di: bool = True) -> ExtractedContent:
        """
        ファイルからテキストを抽出

        ファイル形式に応じて適切な抽出方法を選択します。
        PDF/画像の場合、Document Intelligence が優先されます。

        Args:
            file_name (str): ファイル名
            extension (str): 拡張子（例: ".pdf"）
            base64_content (str): Base64エンコードされたファイル内容
            mime_type (str, optional): MIMEタイプ
            use_di (bool): Document Intelligence を使用するか（デフォルト: True）

        Returns:
            ExtractedContent: 抽出結果

        Note:
            エラーが発生した場合も ExtractedContent が返されます。
            error フィールドでエラー有無を確認できます。
        """
        ext = extension.lower()

        logger.info(f"[抽出開始] {file_name} (形式: {ext})")

        try:
            # PDF・画像の場合、OCRFactoryを使用
            if use_di and ext in (cls.PDF_TYPES + cls.IMAGE_TYPES):
                ocr_result = cls._extract_with_ocr(file_name, base64_content, mime_type)
                if ocr_result:
                    return ocr_result
                # OCRが利用不可の場合はフォールバック

            # ローカル抽出にフォールバック
            if ext in cls.TEXT_TYPES:
                return cls._extract_from_text(file_name, ext, base64_content)

            elif ext in cls.PDF_TYPES:
                return cls._extract_from_pdf(file_name, base64_content)

            elif ext in cls.EXCEL_TYPES:
                return cls._extract_from_excel(file_name, ext, base64_content)

            elif ext in cls.WORD_TYPES:
                return cls._extract_from_word(file_name, ext, base64_content)

            elif ext in cls.IMAGE_TYPES:
                # 画像（Document Intelligence なし）
                logger.warning(f"画像ファイルは Document Intelligence が必要: {file_name}")
                return ExtractedContent(
                    file_name=file_name,
                    file_type="image",
                    text_content=f"[画像ファイル: {file_name}] - Document Intelligence または Vision LLM で処理が必要",
                    extraction_method="metadata_only"
                )

            else:
                logger.warning(f"未対応のファイル形式: {ext}")
                return ExtractedContent(
                    file_name=file_name,
                    file_type="unknown",
                    text_content=f"[未対応ファイル形式: {file_name}]",
                    extraction_method="none",
                    error=f"未対応のファイル形式: {ext}"
                )

        except Exception as e:
            logger.error(f"抽出エラー ({file_name}): {e}")
            return ExtractedContent(
                file_name=file_name,
                file_type=ext,
                text_content=f"[ファイル読み取りエラー: {file_name}]",
                extraction_method="error",
                error=str(e)
            )

    @classmethod
    def _extract_with_ocr(cls, file_name: str, base64_content: str, mime_type: str = None) -> Optional[ExtractedContent]:
        """
        OCRFactoryを使用してテキストを抽出

        OCR_PROVIDER環境変数に基づいて適切なOCRエンジンを使用します。
        OCRが設定されていない場合はNoneを返します。

        Args:
            file_name: ファイル名
            base64_content: Base64エンコードされたファイル内容
            mime_type: MIMEタイプ

        Returns:
            ExtractedContent: 抽出結果（OCR未設定時はNone）
        """
        try:
            from infrastructure.ocr_factory import OCRFactory, OCRProvider

            # OCRプロバイダーを取得
            provider = OCRFactory.get_provider()

            # NONE の場合は None を返してフォールバック
            if provider == OCRProvider.NONE:
                logger.info(f"[OCR] OCR_PROVIDER=NONE, ローカル抽出にフォールバック: {file_name}")
                return None

            # OCRクライアントを取得
            ocr_client = OCRFactory.get_ocr_client()
            if not ocr_client or not ocr_client.is_configured():
                logger.warning(f"[OCR] OCRクライアント未設定, ローカル抽出にフォールバック: {file_name}")
                return None

            logger.info(f"[OCR] {ocr_client.provider_name} を使用: {file_name}")

            # Base64デコード
            file_bytes = base64.b64decode(base64_content)

            # OCR実行
            result = ocr_client.extract_text(file_bytes, mime_type)

            if result.error:
                logger.error(f"[OCR] エラー: {result.error}")
                return ExtractedContent(
                    file_name=file_name,
                    file_type="pdf" if file_name.lower().endswith('.pdf') else "image",
                    text_content=f"[OCRエラー: {file_name}] - {result.error}",
                    extraction_method=f"ocr_error_{provider.value.lower()}",
                    error=result.error
                )

            # OCRResultをExtractedContentに変換
            # 座標情報付きのelementsを変換
            elements = None
            tables = None

            if result.elements:
                elements = [
                    TextElement(
                        element_id=f"elem_{i}",
                        text=elem.text,
                        page_number=elem.page_number,
                        bounding_box=elem.bounding_box or [0, 0, 0, 0],
                        element_type=elem.element_type,
                        confidence=elem.confidence
                    )
                    for i, elem in enumerate(result.elements)
                ]

            if result.tables:
                tables = [
                    ExtractedTable(
                        table_id=table.table_id,
                        page_number=table.page_number,
                        row_count=table.row_count,
                        column_count=table.column_count,
                        cells=[
                            TableCell(
                                row_index=cell.row_index,
                                column_index=cell.column_index,
                                text=cell.text,
                                bounding_box=[0, 0, 0, 0],
                                row_span=cell.row_span,
                                column_span=cell.column_span
                            )
                            for cell in table.cells
                        ],
                        bounding_box=[0, 0, 0, 0]
                    )
                    for table in result.tables
                ]

            return ExtractedContent(
                file_name=file_name,
                file_type="pdf" if file_name.lower().endswith('.pdf') else "image",
                text_content=result.text_content,
                extraction_method=f"ocr_{provider.value.lower()}",
                page_count=result.page_count,
                elements=elements,
                tables=tables
            )

        except ImportError:
            # OCRFactoryがインポートできない場合は後方互換のためDIを試行
            logger.debug("OCRFactory 未使用, AzureDocumentIntelligenceにフォールバック")
            di_client = cls.get_di_client()
            if di_client.is_configured():
                return di_client.extract_with_layout(file_name, base64_content, mime_type)
            return None

        except Exception as e:
            logger.error(f"[OCR] 予期せぬエラー: {e}")
            return None

    @classmethod
    def _extract_from_text(cls, file_name: str, ext: str, base64_content: str) -> ExtractedContent:
        """
        テキストファイルからテキストを抽出

        UTF-8、Shift-JIS のエンコーディングに対応しています。

        Args:
            file_name: ファイル名
            ext: 拡張子
            base64_content: Base64エンコードされた内容

        Returns:
            ExtractedContent: 抽出結果
        """
        try:
            # まず UTF-8 で試行
            content = base64.b64decode(base64_content).decode('utf-8')
            logger.debug(f"UTF-8 でデコード成功: {file_name}")

            return ExtractedContent(
                file_name=file_name,
                file_type="text",
                text_content=content,
                extraction_method="direct_decode"
            )

        except UnicodeDecodeError:
            # Shift-JIS で再試行
            try:
                content = base64.b64decode(base64_content).decode('shift-jis')
                logger.debug(f"Shift-JIS でデコード成功: {file_name}")

                return ExtractedContent(
                    file_name=file_name,
                    file_type="text",
                    text_content=content,
                    extraction_method="direct_decode_shiftjis"
                )

            except Exception as e:
                raise Exception(f"テキストファイルのデコードに失敗: {e}")

    @classmethod
    def _extract_from_pdf(cls, file_name: str, base64_content: str) -> ExtractedContent:
        """
        PDFファイルからテキストを抽出

        pypdf ライブラリを使用します。
        スキャンPDFの場合は Document Intelligence にフォールバックします。

        Args:
            file_name: ファイル名
            base64_content: Base64エンコードされた内容

        Returns:
            ExtractedContent: 抽出結果
        """
        try:
            from pypdf import PdfReader

            # Base64 デコード
            pdf_bytes = base64.b64decode(base64_content)
            pdf_stream = io.BytesIO(pdf_bytes)

            # PDF を読み込み
            reader = PdfReader(pdf_stream)
            page_count = len(reader.pages)
            logger.info(f"PDF読み込み完了: {file_name} ({page_count}ページ)")

            # 各ページからテキストを抽出
            text_parts = []
            has_text = False

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"--- ページ {i+1} ---\n{page_text}")
                    has_text = True

            if has_text:
                full_text = "\n\n".join(text_parts)
                logger.info(f"テキスト抽出成功: {len(full_text):,}文字")

                return ExtractedContent(
                    file_name=file_name,
                    file_type="pdf",
                    text_content=full_text,
                    extraction_method="pypdf",
                    page_count=page_count
                )

            else:
                # スキャンPDF（テキストなし）
                logger.warning(f"スキャンPDFを検出: {file_name}")

                # Document Intelligence でリトライ
                di_client = cls.get_di_client()
                if di_client.is_configured():
                    logger.info(f"Document Intelligence でリトライ: {file_name}")
                    return di_client.extract_with_layout(file_name, base64_content)

                return ExtractedContent(
                    file_name=file_name,
                    file_type="pdf_image",
                    text_content=(
                        f"[スキャンPDF: {file_name}] - {page_count}ページ。"
                        "テキスト抽出不可。Document Intelligence が必要です。"
                    ),
                    extraction_method="pypdf_no_text",
                    page_count=page_count
                )

        except ImportError:
            logger.warning("pypdf がインストールされていません")
            return ExtractedContent(
                file_name=file_name,
                file_type="pdf",
                text_content=f"[PDFファイル: {file_name}] - pypdf ライブラリ未インストール",
                extraction_method="metadata_only",
                error="pypdf がインストールされていません"
            )

        except Exception as e:
            logger.error(f"PDF抽出エラー ({file_name}): {e}")
            return ExtractedContent(
                file_name=file_name,
                file_type="pdf",
                text_content=f"[PDF読み取りエラー: {file_name}] - {str(e)}",
                extraction_method="error",
                error=str(e)
            )

    @classmethod
    def _extract_from_excel(cls, file_name: str, ext: str, base64_content: str) -> ExtractedContent:
        """
        Excelファイルからテキストを抽出

        openpyxl ライブラリを使用します。
        全シートの内容をタブ区切りテキストとして抽出します。

        Args:
            file_name: ファイル名
            ext: 拡張子
            base64_content: Base64エンコードされた内容

        Returns:
            ExtractedContent: 抽出結果
        """
        try:
            from openpyxl import load_workbook

            # Base64 デコード
            excel_bytes = base64.b64decode(base64_content)
            excel_stream = io.BytesIO(excel_bytes)

            # Excel を読み込み（数式は計算済みの値を取得）
            wb = load_workbook(excel_stream, data_only=True)
            logger.info(f"Excel読み込み完了: {file_name} ({len(wb.sheetnames)}シート)")

            text_parts = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_text = [f"=== シート: {sheet_name} ==="]

                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_values):  # 空行はスキップ
                        sheet_text.append("\t".join(row_values))

                text_parts.append("\n".join(sheet_text))

            full_text = "\n\n".join(text_parts)
            logger.info(f"Excel抽出成功: {len(full_text):,}文字")

            return ExtractedContent(
                file_name=file_name,
                file_type="excel",
                text_content=full_text,
                extraction_method="openpyxl",
                page_count=len(wb.sheetnames)
            )

        except ImportError:
            logger.warning("openpyxl がインストールされていません")
            return ExtractedContent(
                file_name=file_name,
                file_type="excel",
                text_content=f"[Excelファイル: {file_name}] - openpyxl ライブラリ未インストール",
                extraction_method="metadata_only",
                error="openpyxl がインストールされていません"
            )

        except Exception as e:
            logger.error(f"Excel抽出エラー ({file_name}): {e}")
            return ExtractedContent(
                file_name=file_name,
                file_type="excel",
                text_content=f"[Excel読み取りエラー: {file_name}] - {str(e)}",
                extraction_method="error",
                error=str(e)
            )

    @classmethod
    def _extract_from_word(cls, file_name: str, ext: str, base64_content: str) -> ExtractedContent:
        """
        Wordファイル（.docx）からテキストを抽出

        python-docx ライブラリを使用します。
        段落、表、ヘッダー、フッターからテキストを抽出します。

        Args:
            file_name: ファイル名
            ext: 拡張子
            base64_content: Base64エンコードされた内容

        Returns:
            ExtractedContent: 抽出結果
        """
        try:
            from docx import Document

            # Base64 デコード
            word_bytes = base64.b64decode(base64_content)
            word_stream = io.BytesIO(word_bytes)

            # Word を読み込み
            doc = Document(word_stream)
            logger.info(f"Word読み込み完了: {file_name}")

            text_parts = []

            # 段落を抽出
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 表を抽出
            for table_idx, table in enumerate(doc.tables):
                text_parts.append(f"\n=== 表 {table_idx + 1} ===")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    if any(row_text):
                        text_parts.append("\t".join(row_text))

            full_text = "\n".join(text_parts)
            logger.info(f"Word抽出成功: {len(full_text):,}文字")

            return ExtractedContent(
                file_name=file_name,
                file_type="word",
                text_content=full_text,
                extraction_method="python-docx",
                page_count=None  # DOCXではページ数の取得が困難
            )

        except ImportError:
            logger.warning("python-docx がインストールされていません")
            return ExtractedContent(
                file_name=file_name,
                file_type="word",
                text_content=f"[Wordファイル: {file_name}] - python-docx ライブラリ未インストール",
                extraction_method="metadata_only",
                error="python-docx がインストールされていません"
            )

        except Exception as e:
            logger.error(f"Word抽出エラー ({file_name}): {e}")
            return ExtractedContent(
                file_name=file_name,
                file_type="word",
                text_content=f"[Word読み取りエラー: {file_name}] - {str(e)}",
                extraction_method="error",
                error=str(e)
            )

    @classmethod
    def extract_all(cls, evidence_files: List, use_di: bool = True) -> List[ExtractedContent]:
        """
        複数の証跡ファイルを一括抽出

        Args:
            evidence_files: EvidenceFile オブジェクトのリスト
            use_di: Document Intelligence を使用するか

        Returns:
            List[ExtractedContent]: 抽出結果のリスト
        """
        logger.info(f"=== 一括抽出開始: {len(evidence_files)}ファイル ===")

        results = []
        for i, ef in enumerate(evidence_files, 1):
            logger.info(f"[{i}/{len(evidence_files)}] {ef.file_name}")

            result = cls.extract_text(
                file_name=ef.file_name,
                extension=ef.extension,
                base64_content=ef.base64_content,
                mime_type=ef.mime_type,
                use_di=use_di
            )
            results.append(result)

            # 結果サマリをログ出力
            status = "成功" if not result.error else f"エラー: {result.error}"
            logger.debug(f"  → {result.extraction_method}: {status}")

        logger.info(f"=== 一括抽出完了: {len(results)}ファイル処理 ===")
        return results

    @classmethod
    def format_for_prompt(cls, extracted_contents: List[ExtractedContent],
                          max_chars_per_file: int = 10000) -> str:
        """
        LLMプロンプト用にテキストをフォーマット

        抽出したテキストをLLMに渡す形式に整形します。
        長いテキストは自動的に切り詰められます。

        Args:
            extracted_contents: 抽出結果のリスト
            max_chars_per_file: ファイルあたりの最大文字数

        Returns:
            str: フォーマットされたテキスト

        Example:
            >>> text = DocumentProcessor.format_for_prompt(results)
            >>> # 結果例:
            >>> # 【ファイル: report.pdf】 (3ページ) [抽出方法: pypdf]
            >>> # --- ページ 1 ---
            >>> # ...
        """
        parts = []

        for ec in extracted_contents:
            content = ec.text_content

            # 長いテキストは切り詰め
            if len(content) > max_chars_per_file:
                content = content[:max_chars_per_file] + \
                    f"\n... (以下省略、全{len(ec.text_content):,}文字)"

            # ヘッダー作成
            header = f"【ファイル: {ec.file_name}】"
            if ec.page_count:
                header += f" ({ec.page_count}ページ)"
            header += f" [抽出方法: {ec.extraction_method}]"

            parts.append(f"{header}\n{content}")

        if parts:
            return "\n\n" + "="*50 + "\n\n".join(parts)
        else:
            return "エビデンスデータなし"

    @classmethod
    def format_for_prompt_with_elements(cls, extracted_contents: List[ExtractedContent],
                                        max_chars_per_file: int = 10000) -> str:
        """
        要素ID付きでLLMプロンプト用にフォーマット

        将来的なハイライト機能用に、要素IDを含むフォーマットを生成します。

        Args:
            extracted_contents: 抽出結果のリスト
            max_chars_per_file: ファイルあたりの最大文字数

        Returns:
            str: 要素ID付きのフォーマットされたテキスト
        """
        parts = []

        for ec in extracted_contents:
            header = f"【ファイル: {ec.file_name}】"
            if ec.page_count:
                header += f" ({ec.page_count}ページ)"

            if ec.elements:
                # 要素IDを含むフォーマット
                content_parts = []
                for elem in ec.elements:
                    if elem.element_type in ["paragraph", "line"]:
                        content_parts.append(f"[{elem.element_id}] {elem.text}")
                content = "\n".join(content_parts)
            else:
                content = ec.text_content

            if len(content) > max_chars_per_file:
                content = content[:max_chars_per_file] + "\n... (以下省略)"

            parts.append(f"{header}\n{content}")

        return "\n\n".join(parts) if parts else "エビデンスデータなし"

    @classmethod
    def get_element_by_id(cls, extracted_contents: List[ExtractedContent],
                          element_id: str) -> Optional[TextElement]:
        """
        要素IDから要素を検索（ハイライト機能用）

        Args:
            extracted_contents: 抽出結果のリスト
            element_id: 検索する要素ID

        Returns:
            TextElement: 見つかった場合は要素、見つからない場合は None
        """
        for ec in extracted_contents:
            if ec.elements:
                for elem in ec.elements:
                    if elem.element_id == element_id:
                        return elem
        return None

    @classmethod
    def get_config_status(cls) -> Dict[str, Any]:
        """
        Document Processor の設定状態を取得

        Returns:
            dict: 設定状態
                - ocr_provider: OCRプロバイダー名
                - ocr_configured: OCR設定済みか
                - document_intelligence_configured: DI設定済みか（後方互換）
                - supported_formats: 対応ファイル形式
        """
        # OCRFactory経由の設定状態
        ocr_status = {
            "provider": "NONE",
            "configured": False,
            "provider_name": "OCR無効"
        }

        try:
            from infrastructure.ocr_factory import OCRFactory
            ocr_config = OCRFactory.get_config_status()
            ocr_status["provider"] = ocr_config.get("provider", "NONE")
            ocr_status["configured"] = ocr_config.get("configured", False)

            ocr_client = OCRFactory.get_ocr_client()
            if ocr_client:
                ocr_status["provider_name"] = ocr_client.provider_name
        except ImportError:
            pass

        # 後方互換: AzureDocumentIntelligenceの状態
        di_client = cls.get_di_client()
        di_configured = di_client.is_configured()

        return {
            "ocr_provider": ocr_status["provider"],
            "ocr_provider_name": ocr_status["provider_name"],
            "ocr_configured": ocr_status["configured"],
            # 後方互換フィールド
            "document_intelligence_configured": di_configured or (ocr_status["provider"] == "AZURE" and ocr_status["configured"]),
            "endpoint": di_client.endpoint[:30] + "..." if di_client.endpoint else None,
            "supported_formats": {
                "text": cls.TEXT_TYPES,
                "pdf": cls.PDF_TYPES,
                "excel": cls.EXCEL_TYPES,
                "word": cls.WORD_TYPES,
                "image": cls.IMAGE_TYPES
            }
        }
